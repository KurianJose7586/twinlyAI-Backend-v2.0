# app/core/rag_pipeline.py

import asyncio
import json
import logging
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_mongodb import MongoDBAtlasVectorSearch
from pymongo import MongoClient
import certifi
from langchain_huggingface import HuggingFaceEmbeddings
# Using local HuggingFaceEmbeddings — model is pre-baked into Docker image during build
# (see Dockerfile RUN step), so there is no runtime download delay and no dependency
# on HuggingFace Inference API tokens or permissions.
from langchain_groq import ChatGroq
from langgraph.prebuilt import create_react_agent
from langchain_core.tools import tool
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, AIMessage
from app.core.config import settings

from langchain_core.output_parsers import JsonOutputParser
from pydantic import BaseModel, Field
from typing import List

# --- GLOBAL MODEL CACHE, DB CLIENT & LLM ---
_EMBEDDINGS_MODEL = None
_SYNC_MONGO_CLIENT = None
_LLM_INSTANCE = None
_EXTRACTION_LLM = None            # ChatGroq for metadata/interview extraction
_BOT_VECTOR_STORE = None          # MongoDBAtlasVectorStore for vector_store_bots
_GLOBAL_VECTOR_STORE = None        # MongoDBAtlasVectorStore for vector_store_global


def get_embeddings_model():
    global _EMBEDDINGS_MODEL
    if _EMBEDDINGS_MODEL is None:
        logging.info("Loading local HuggingFaceEmbeddings (model pre-baked in Docker image)...")
        _EMBEDDINGS_MODEL = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
        )
    return _EMBEDDINGS_MODEL


def get_sync_mongo_client():
    global _SYNC_MONGO_CLIENT
    if _SYNC_MONGO_CLIENT is None:
        client_kwargs = {}
        if "mongodb+srv://" in settings.MONGO_CONNECTION_STRING:
            client_kwargs["tlsCAFile"] = certifi.where()
        _SYNC_MONGO_CLIENT = MongoClient(settings.MONGO_CONNECTION_STRING, **client_kwargs)
    return _SYNC_MONGO_CLIENT


def get_llm():
    """Module-level singleton for the primary LLM (used in chat streaming)."""
    global _LLM_INSTANCE
    if _LLM_INSTANCE is None:
        _LLM_INSTANCE = ChatGroq(
            model_name="meta-llama/llama-4-scout-17b-16e-instruct",
            temperature=0.7,
            groq_api_key=settings.GROQ_API_KEY,
        )
    return _LLM_INSTANCE


def _get_bot_vector_store():
    """Module-level singleton for the bot-specific vector store."""
    global _BOT_VECTOR_STORE
    if _BOT_VECTOR_STORE is None:
        client = get_sync_mongo_client()
        db = client[settings.MONGO_DB_NAME]
        _BOT_VECTOR_STORE = MongoDBAtlasVectorSearch(
            collection=db["vector_store_bots"],
            embedding=get_embeddings_model(),
            index_name="vector_index",
            text_key="text",
            embedding_key="embedding",
        )
    return _BOT_VECTOR_STORE


def _get_global_vector_store():
    """Module-level singleton for the global recruiter index vector store."""
    global _GLOBAL_VECTOR_STORE
    if _GLOBAL_VECTOR_STORE is None:
        client = get_sync_mongo_client()
        db = client[settings.MONGO_DB_NAME]
        _GLOBAL_VECTOR_STORE = MongoDBAtlasVectorSearch(
            collection=db["vector_store_global"],
            embedding=get_embeddings_model(),
            index_name="vector_index",
            text_key="text",
            embedding_key="embedding",
        )
    return _GLOBAL_VECTOR_STORE


def get_extraction_llm():
    """Module-level singleton for the extraction LLM (metadata & interview analysis)."""
    global _EXTRACTION_LLM
    if _EXTRACTION_LLM is None:
        _EXTRACTION_LLM = ChatGroq(
            model_name="llama-3.3-70b-versatile",
            temperature=0.0,
            groq_api_key=settings.GROQ_API_KEY,
        )
    return _EXTRACTION_LLM


# --- Pydantic model for metadata extraction ---
class ResumeMetadata(BaseModel):
    candidate_name: str = Field(description="The full name of the candidate")
    summary: str = Field(description="A concise 2-sentence professional summary of the candidate")
    skills: List[str] = Field(description="A list of the top 10 most relevant technical skills or tools")
    experience_years: float = Field(description="Total estimated years of professional experience as a number (e.g., 3.5)")

class InterviewAssessment(BaseModel):
    topics_covered: List[str] = Field(description="High-level topics or skills discussed so far")
    red_flags: List[str] = Field(description="Any concerns or missing skills that surfaced during the chat")
    recruiter_intent: str = Field(description="What the recruiter seems to be looking for or prioritizing")


# --- JSON to TEXT CONVERSION (Helper Function) ---
def json_to_text(json_data: dict) -> str:
    text = ""
    for key, value in json_data.items():
        if isinstance(value, dict):
            text += "{}:\n".format(key.replace('_', ' ').title())
            for sub_key, sub_value in value.items():
                text += "  {}: {}\n".format(sub_key.replace('_', ' ').title(), sub_value)
        elif isinstance(value, list):
            text += "{}:\n".format(key.replace('_', ' ').title())
            for item in value:
                if isinstance(item, dict):
                    for item_key, item_value in item.items():
                        text += "  - {}: {}\n".format(item_key.replace('_', ' ').title(), item_value)
                else:
                    text += "- {}\n".format(item)
        else:
            text += "{}: {}\n".format(key.replace('_', ' ').title(), value)
    return text


# --- FILE PROCESSING (Helper Function) ---
def extract_text_from_file(file_path: Path) -> str:
    text = ""
    try:
        if file_path.suffix.lower() == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                text = "".join(page.extract_text() or "" for page in pdf.pages)
        elif file_path.suffix.lower() == ".docx":
            doc = DocxDocument(file_path)
            text = "\n".join(para.text for para in doc.paragraphs)
        elif file_path.suffix.lower() == ".txt":
            text = file_path.read_text(encoding="utf-8")
        elif file_path.suffix.lower() == ".json":
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                text = json_to_text(data)
        else:
            raise ValueError("Unsupported file type: {}".format(file_path.suffix))
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {e}")
        raise ValueError(f"Could not read file content: {str(e)}")

    if not text.strip():
        raise ValueError("The uploaded file contains no extractable text. If this is a PDF, ensure it is not an image scan.")

    return text


class RAGPipeline:
    def __init__(self, bot_id: str, user_id: str, bot_name: str, user_email: str = ""):
        self.bot_id = bot_id
        self.user_id = user_id
        self.bot_name = bot_name
        self.user_email = user_email

        # Use module-level singletons to avoid per-request connection overhead
        self.embeddings = get_embeddings_model()
        self.vector_store = _get_bot_vector_store()
        self.llm = get_llm()

        # Direct DB handle for count_documents / delete_many (lightweight)
        self.mongo_client = get_sync_mongo_client()
        self.db = self.mongo_client[settings.MONGO_DB_NAME]
        self.collection = self.db["vector_store_bots"]

        self.github_vector_store = None  # Loaded lazily on first use
        self._github_checked = False
        self.agent_executor = None

    def _ensure_github_checked(self):
        """Lazily check GitHub vector store status (only once)."""
        if not self._github_checked:
            if self.user_email:
                try:
                    count = self.db["vector_store_github"].count_documents({"user_email": self.user_email})
                    self.github_vector_store = True if count > 0 else None
                except Exception:
                    self.github_vector_store = None
            self._github_checked = True

    def _create_agent(self, dynamic_metadata_text: str = ""):
        self._ensure_github_checked()
        has_github = self.github_vector_store is not None

        # Filter for bot-specific documents only
        retriever = self.vector_store.as_retriever(
            search_kwargs={"k": 5, "pre_filter": {"bot_id": self.bot_id}}
        )

        @tool
        def search_resume(query: str) -> str:
            """Search the candidate's resume for career details, skills, and history."""
            try:
                docs = retriever.invoke(query)
                return "\n\n".join([d.page_content for d in docs]) or "No relevant info found."
            except Exception as e:
                logging.error(f"Search resume tool error: {e}")
                return "Resume search is temporarily unavailable."

        @tool
        def search_github_code(query: str) -> str:
            """Search the candidate's synced GitHub repositories for code examples."""
            try:
                gh_store = MongoDBAtlasVectorSearch(
                    collection=self.db["vector_store_github"],
                    embedding=self.embeddings,
                    index_name="vector_index"
                )
                docs = gh_store.similarity_search(
                    query, k=3, pre_filter={"user_email": self.user_email}
                )
                return "\n\n".join([f"[{d.metadata.get('path')}]\n{d.page_content}" for d in docs]) or "No code found."
            except Exception as e:
                logging.error(f"GH search error: {e}")
                return "GitHub search unavailable."

        @tool
        def calculate_experience(start_year: float, end_year: float) -> float:
            """Compute the number of years between two years."""
            return max(0, end_year - start_year)

        tools = [search_resume, calculate_experience]
        if has_github: tools.append(search_github_code)

        system_prompt = (
            'You are "{bot_name}", a professional AI Twin.\n'
            'Candidate Profile:\n---\n{metadata}\n---\n'
            'Rules: Always use tools for facts. Speak in third person. Be accurate.'
        ).format(bot_name=self.bot_name, metadata=dynamic_metadata_text)

        return create_react_agent(model=self.llm, tools=tools, prompt=system_prompt)

    async def process_file(self, file_path: str):
        try:
            text_content = extract_text_from_file(Path(file_path))
            if not text_content: return False

            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            splits = text_splitter.split_text(text_content)

            docs = [
                Document(page_content=s, metadata={"bot_id": self.bot_id, "source": Path(file_path).name})
                for s in splits
            ]

            # Clear old vectors for this bot — run blocking call in thread
            await asyncio.to_thread(self.collection.delete_many, {"bot_id": self.bot_id})

            # Store new ones — run blocking call in thread
            await asyncio.to_thread(self.vector_store.add_documents, docs)
            return True
        except Exception:
            logging.exception("Error in process_file")
            raise

    async def extract_metadata(self, file_path: str) -> dict:
        """Uses LLM to extract skills and summary."""
        text_content = extract_text_from_file(Path(file_path))
        parser = JsonOutputParser(pydantic_object=ResumeMetadata)
        prompt = ChatPromptTemplate.from_messages([
            ("system", "Extract JSON metadata from this resume text.\n{format_instructions}"),
            ("human", "{resume_text}")
        ])
        chain = prompt | get_extraction_llm() | parser
        try:
            return await chain.ainvoke({"resume_text": text_content[:10000], "format_instructions": parser.get_format_instructions()})
        except Exception:
            return {"candidate_name": self.bot_name, "summary": "N/A", "skills": [], "experience_years": 0}

    async def analyze_interview(self, chat_history: list) -> dict:
        """Analyzes state of the interview."""
        if len(chat_history) < 2: return {}
        history_text = "\n".join(["{}: {}".format(getattr(msg, 'type', 'unknown'), msg.content) for msg in chat_history[-10:]])
        parser = JsonOutputParser(pydantic_object=InterviewAssessment)
        prompt = ChatPromptTemplate.from_messages([
            ("system", "Analyze this interview chat. Return JSON.\n{format_instructions}"),
            ("human", "{history_text}")
        ])
        chain = prompt | get_extraction_llm() | parser
        try:
            return await chain.ainvoke({
                "history_text": history_text,
                "format_instructions": parser.get_format_instructions()
            })
        except Exception: return {}

    async def get_response_stream(self, user_message: str, chat_history: list = [], bot_metadata: dict = None):
        meta_context = ""
        if bot_metadata:
            meta_context = "Name: {}\nSummary: {}\nSkills: {}".format(
                bot_metadata.get('name', self.bot_name),
                bot_metadata.get('summary', 'N/A'),
                ', '.join(bot_metadata.get('skills') or [])
            )

        if not user_message.strip():
            yield "Please say something."
            return

        # Ensure GitHub status is checked
        self._ensure_github_checked()

        # Check if indexed docs exist to use the Agent — run blocking call in thread
        count = await asyncio.to_thread(self.collection.count_documents, {"bot_id": self.bot_id})
        if count > 0 or self.github_vector_store:
            self.agent_executor = self._create_agent(dynamic_metadata_text=meta_context)

        messages = chat_history + [HumanMessage(content=user_message)]

        if not self.agent_executor:
            # Fallback to pure LLM with metadata
            fallback_prompt = ChatPromptTemplate.from_messages([
                ("system", "You are {bot_name}. Answer based on: {meta}"),
                MessagesPlaceholder(variable_name="history"),
                ("human", "{input}")
            ])
            chain = fallback_prompt | self.llm
            async for chunk in chain.astream({"input": user_message, "history": chat_history, "bot_name": self.bot_name, "meta": meta_context}):
                if hasattr(chunk, "content"): yield chunk.content
        else:
            async for event in self.agent_executor.astream_events({"messages": messages}, version="v2"):
                if event["event"] == "on_chat_model_stream":
                    content = event["data"]["chunk"].content
                    if content: yield content


# --- GLOBAL RECRUITER INDEX ---
class GlobalRecruiterIndex:
    def __init__(self):
        # Use module-level singletons — same MongoClient & embeddings as RAGPipeline
        self.vector_store = _get_global_vector_store()
        self.mongo_client = get_sync_mongo_client()
        self.collection = self.mongo_client[settings.MONGO_DB_NAME]["vector_store_global"]

    async def add_candidate_profile(self, bot_id: str, profile_text: str):
        doc = Document(page_content=profile_text, metadata={"bot_id": bot_id})
        # Run blocking calls in threads
        await asyncio.to_thread(self.collection.delete_many, {"bot_id": bot_id})
        await asyncio.to_thread(self.vector_store.add_documents, [doc])
        return True

    async def semantic_search(self, query: str, k: int = 10) -> List[str]:
        try:
            results = await asyncio.to_thread(self.vector_store.similarity_search, query, k=k)
            return [str(d.metadata.get("bot_id")) for d in results if d.metadata.get("bot_id")]
        except Exception as e:
            logging.error(f"Global search error: {e}")
            raise RuntimeError("Search failed.")
